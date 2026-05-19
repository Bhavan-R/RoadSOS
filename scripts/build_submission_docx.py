"""Build the National Road Safety Hackathon 2026 submission Word document.

The CoERS x IIT Madras rulebook requires participants to share:
  (1) The entire code
  (2) A list of software packages used
  (3) Assumptions

This script bundles all three into a single .docx that judges can open in Word.

Usage:
  python scripts/build_submission_docx.py
Output:
  docs/RoadSOS_Submission.docx
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parents[1]


# ──────────────────────────────────────────────────────────────────────────────
# Files to include in the code appendix
# ──────────────────────────────────────────────────────────────────────────────

CODE_GROUPS: list[tuple[str, list[str]]] = [
    (
        "Root configuration",
        [
            "README.md",
            "CLAUDE.md",
            "vercel.json",
            ".github/workflows/backend-tests.yml",
            ".github/workflows/frontend-ci.yml",
            ".github/workflows/pr-guard.yml",
        ],
    ),
    (
        "Backend - top level",
        [
            "backend/requirements.txt",
            "backend/requirements-dev.txt",
            "backend/ruff.toml",
            "backend/main.py",
            "backend/middleware.py",
            "backend/logging_config.py",
        ],
    ),
    (
        "Backend - services",
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "backend" / "services").glob("*.py")),
    ),
    (
        "Backend - tests",
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "backend" / "tests").glob("*.py")),
    ),
    (
        "Frontend - configuration",
        [
            "frontend/package.json",
            "frontend/vite.config.js",
            "frontend/index.html",
            "frontend/vercel.json",
            "frontend/tsconfig.json",
        ],
    ),
    (
        "Frontend - entry point",
        ["frontend/src/App.jsx", "frontend/src/constants.js"],
    ),
    (
        "Frontend - components",
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "frontend" / "src" / "components").glob("*.jsx")),
    ),
    (
        "Frontend - hooks",
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "frontend" / "src" / "hooks").glob("*.js")),
    ),
    (
        "Frontend - utilities",
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "frontend" / "src" / "utils").glob("*.js")),
    ),
    (
        "Frontend - styles",
        ["frontend/src/style.css", "frontend/src/final-design.css"],
    ),
    (
        "Frontend - bundled data",
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "frontend" / "src" / "data").glob("*.json")),
    ),
    (
        "Frontend - i18n (48 locale bundles)",
        ["frontend/src/i18n/index.js", "frontend/src/i18n/locales.js",
         "frontend/src/i18n/countryLanguage.js"] +
        sorted(str(p.relative_to(ROOT)).replace("\\", "/")
               for p in (ROOT / "frontend" / "src" / "i18n").glob("*.json")),
    ),
    (
        "Frontend - tests",
        sorted([str(p.relative_to(ROOT)).replace("\\", "/")
                for p in (ROOT / "frontend" / "src" / "__tests__").glob("*.js")] +
               [str(p.relative_to(ROOT)).replace("\\", "/")
                for p in (ROOT / "frontend" / "src" / "utils" / "__tests__").glob("*.js")]),
    ),
]

# Language extension → Word "language" for syntax hint (we only use monospace).
LANG_BY_SUFFIX = {
    ".py": "Python", ".js": "JavaScript", ".jsx": "JSX",
    ".json": "JSON", ".css": "CSS", ".html": "HTML",
    ".yml": "YAML", ".yaml": "YAML", ".toml": "TOML",
    ".md": "Markdown", ".txt": "Text",
}


# ──────────────────────────────────────────────────────────────────────────────
# Docx helpers
# ──────────────────────────────────────────────────────────────────────────────

def _set_cell_shading(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x14, 0x24)


def add_para(doc: Document, text: str, *, bold: bool = False,
             italic: bool = False, size_pt: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)


def add_code_block(doc: Document, text: str, *, max_chars: int = 200_000) -> None:
    """Insert a monospaced code block.

    Word handles very large single paragraphs poorly, so we split per line.
    Files longer than max_chars are truncated with a marker.
    """
    truncated = False
    if len(text) > max_chars:
        text = text[:max_chars]
        truncated = True

    # Use a 1-column table with grey shading to render a code block.
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "F1F5F9")

    # First paragraph already exists in the cell; reuse it.
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)

    lines = text.splitlines() or [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = para
        else:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        # Set the East Asian font too so Word doesn't substitute.
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Consolas")
        rFonts.set(qn("w:hAnsi"), "Consolas")
        rFonts.set(qn("w:cs"), "Consolas")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    if truncated:
        p = cell.add_paragraph()
        r = p.add_run(f"  … (file truncated at {max_chars:,} characters) …")
        r.italic = True
        r.font.size = Pt(8.5)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ──────────────────────────────────────────────────────────────────────────────
# Content
# ──────────────────────────────────────────────────────────────────────────────

ASSUMPTIONS: list[tuple[str, list[str]]] = [
    ("Network and connectivity", [
        "The user's device has at least intermittent connectivity at the moment of first install (so the service worker can precache the shell and bundled facilities). After install, the app is fully usable offline.",
        "When online, OpenStreetMap's Overpass API and Google Places API are reachable; when not, the four-tier fallback chain (live → localStorage cache → bundled JSON → hardcoded mock) ensures the app still returns contacts.",
        "Render's free-tier backend spins down after 15 minutes of inactivity, so the first request after idle can take up to 55 seconds. The frontend includes a warm-up ping and a status banner explaining this to the user.",
    ]),
    ("Device and browser", [
        "Modern evergreen browsers (Chrome, Edge, Firefox, Safari ≥ 14). Service workers, the Geolocation API, and the Web Speech API are required for full functionality.",
        "Battery Status API is supported on Chromium-based browsers (Chrome, Edge, Opera, Android WebView) but not on iOS Safari or desktop Firefox. The Battery row in the dispatch screen is hidden when the API is unavailable rather than showing a placeholder.",
        "Accelerometer access (used for crash detection) is granted on devices exposing the DeviceMotion API. iOS Safari requires an explicit permission prompt; the app degrades gracefully if denied.",
    ]),
    ("Geolocation", [
        "GPS fixes are accurate to within 50 metres in open sky and 200 metres in urban canyons. The app commits a coarse fix in under 2 s so the user sees something on the map immediately, then refines with watchPosition.",
        "iCloud Private Relay and similar VPN/proxy services can return IP-geolocated locations that are hundreds of kilometres from the real device. A three-layer guard (org-name check, > 500 km haversine drift from last known GPS, Apple datacentre bbox) rejects these and re-prompts for GPS.",
        "Locations are cached at ~110 m grid resolution (3 decimal places) — coarse enough for high cache hit-rate during demos, fine enough to keep emergency contacts locally relevant.",
    ]),
    ("Geography and data coverage", [
        "OpenStreetMap data quality is best in India, Europe, and dense urban areas worldwide. In sparse rural regions, the app expands the Overpass search radius from 5 km → 10 km → 20 km and falls back on bundled facilities (249 facilities across 196 countries).",
        "Country emergency numbers (police, ambulance, fire, disaster) are pre-bundled for all 196 countries and always render without any network call.",
        "ISO-3166 country code is resolved from Nominatim reverse-geocode; if Nominatim fails the app falls back to coarse bounding-box matching against the bundled country dataset.",
    ]),
    ("User behaviour and consent", [
        "Users grant geolocation permission on first launch; otherwise the app shows a 'Set location manually' affordance that lets the user tap on the map or search an address.",
        "Users entering Medical ID data (blood type, allergies, emergency contacts) consent to that data being included in outgoing SOS SMS / WhatsApp payloads. The data is stored only in localStorage on the device — never sent to RoadSOS servers.",
        "When a crash is auto-detected (≥ 25 km/h velocity collapse to ≤ 5 km/h within 2 s, confirmed by an accelerometer spike ≥ 3.5 G), the app shows a 10 s cancellation window before auto-dispatching, allowing the user to abort false positives.",
    ]),
    ("Communications channels", [
        "Country code determines the primary dispatch channel: WhatsApp-dominant countries (India, Brazil, Indonesia, Nigeria, Mexico, Argentina, Pakistan, Bangladesh, etc.) prefer wa.me links; everywhere else defaults to native SMS via sms: deep links.",
        "If the WhatsApp app is not installed, the wa.me link opens WhatsApp Web in a browser tab. The app deferred-falls-back to SMS after 1.2 s if the WA window failed to open.",
        "The window.open call for WhatsApp must fire synchronously inside the click handler — browsers strip the user-gesture flag after any await, so popups are blocked. The handler therefore opens the dispatch channel first and only then awaits camera/alarm side effects.",
    ]),
    ("AI triage", [
        "Google Gemini 2.0 Flash is used for situational triage via direct REST calls (no SDK). Free-tier limits (60 RPM / 1500 RPD) are sufficient for a demo and small pilot.",
        "Triage is independent of search: if Gemini is unreachable or rate-limited, the app falls back to a deterministic rule-based prioritisation that uses the same two yes/no questions (Is anyone injured? Is the vehicle blocking the road?).",
    ]),
    ("Hosting", [
        "Frontend is hosted on Vercel with auto-deploy from main. Backend is hosted on Render with auto-deploy from main. Both are free tiers — production-grade hosting would migrate the backend to a paid Fluid Compute or Cloud Run instance to eliminate cold starts.",
        "The frontend can be served from any static host; no Vercel-specific features are used. The backend is a vanilla FastAPI app and can be redeployed on any Python 3.11+ host that accepts uvicorn.",
    ]),
    ("Security and privacy", [
        "Medical ID, language preference, and cached searches live in localStorage only. The backend is stateless and never persists user data.",
        "Rate limits (/search 30 rpm/IP, /triage 20 rpm/IP) are enforced by a token-bucket middleware that is aware of Render's reverse-proxy X-Forwarded-For header.",
        "Google Places API keys, when provided, are rotated across a comma-separated list on every request; if a key returns 403 (quota) the rotation picks the next key. Keys are server-side only; the frontend never sees them.",
    ]),
]


def _packages_section(doc: Document) -> None:
    add_heading(doc, "Software Packages", level=1)

    add_heading(doc, "Backend — Python runtime (backend/requirements.txt)", level=2)
    backend_reqs = (ROOT / "backend" / "requirements.txt").read_text(encoding="utf-8")
    add_code_block(doc, backend_reqs)

    add_heading(doc, "Backend — Python development tools (backend/requirements-dev.txt)", level=2)
    dev_path = ROOT / "backend" / "requirements-dev.txt"
    if dev_path.exists():
        add_code_block(doc, dev_path.read_text(encoding="utf-8"))
    else:
        add_para(doc, "(file not present in this revision)", italic=True)

    add_heading(doc, "Frontend — npm dependencies (frontend/package.json)", level=2)
    pkg = json.loads((ROOT / "frontend" / "package.json").read_text(encoding="utf-8"))

    def _table(title: str, deps: dict[str, str]) -> None:
        add_para(doc, title, bold=True)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Package"
        hdr[1].text = "Version"
        for name, version in sorted(deps.items()):
            row = table.add_row().cells
            row[0].text = name
            row[1].text = str(version)
        doc.add_paragraph()

    _table("Runtime dependencies", pkg.get("dependencies", {}))
    _table("Development dependencies", pkg.get("devDependencies", {}))

    add_heading(doc, "External services (no SDK / used via REST)", level=2)
    services = [
        ("OpenStreetMap Overpass API", "Public OSM data query layer. Three mirrors with exponential-backoff retry."),
        ("OpenStreetMap Nominatim", "Reverse geocoding (lat/lon → country code, landmark)."),
        ("Google Places API (Nearby Search + Place Details)", "Optional parallel fallback for sparse OSM regions."),
        ("Google Gemini 2.0 Flash", "AI triage. Free tier: 60 RPM / 1 500 RPD."),
        ("CartoDB Dark Matter tiles", "Map tile provider (no API key required)."),
        ("Web platform APIs", "Geolocation, DeviceMotion, Battery Status, Web Speech, Service Worker, IndexedDB (via Workbox)."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Service / API"
    table.rows[0].cells[1].text = "Role in RoadSOS"
    for name, role in services:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = role


def _assumptions_section(doc: Document) -> None:
    add_heading(doc, "Assumptions", level=1)
    add_para(doc,
             "The following assumptions underpin the design and the judging-criteria mapping. "
             "They are the constraints the team has deliberately reasoned about; the code is "
             "engineered to degrade gracefully whenever any single assumption is violated at "
             "runtime.",
             italic=True)
    for heading, bullets in ASSUMPTIONS:
        add_heading(doc, heading, level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")


def _overview_section(doc: Document) -> None:
    add_heading(doc, "Project overview", level=1)
    add_para(doc, "RoadSOS — emergency response and prioritised contact discovery for road accidents.", bold=True)
    add_para(doc,
             "India records 1.5 lakh road-accident deaths per year. The first 60 minutes after "
             "a severe injury (the 'golden hour') determines survival. At a real crash scene, "
             "a bystander currently needs 2–3 minutes just to find the right phone number via "
             "Google Maps. RoadSOS reduces that to under 10 seconds and works fully offline.")

    add_heading(doc, "Submission contents", level=2)
    bullets = [
        "Section 1 — Assumptions: explicit list of constraints under which RoadSOS operates.",
        "Section 2 — Software packages: backend pip dependencies, frontend npm dependencies (runtime + dev), and the external services / web-platform APIs consumed.",
        "Section 3 — Architecture summary: a one-page reference; the full 35-page reference is in docs/TECHNICAL.pdf in the repository.",
        "Section 4 — Entire source code: every file in backend/, frontend/src/, the CI workflows, and the deployment configuration, grouped by directory and rendered verbatim.",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "Judging-criteria mapping (from the CoERS rulebook)", level=2)
    criteria = [
        ("Number of emergency contacts found",
         "OSM Overpass (9 categories in parallel) + Google Places (parallel, not sequential). Auto-expanding radius 5 → 10 → 20 km."),
        ("Reliability",
         "Every external call wrapped in _safe_* helpers that never raise. 3-mirror Overpass with exponential backoff. API always returns HTTP 200 with a valid shape."),
        ("Offline functionality",
         "Four-tier fallback: live backend → 24 h localStorage cache → bundled JSON (249 facilities) → hardcoded mock. Country emergency numbers always render offline."),
        ("Information integration across countries",
         "196 countries pre-loaded. ISO-3166 country code from Nominatim. Emergency numbers switch automatically at borders."),
        ("Six mandatory categories",
         "hospital · police · ambulance · towing · tyre/puncture · showroom — all mapped to OSM tags and Google keyword queries."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.rows[0].cells[0].text = "Criterion"
    table.rows[0].cells[1].text = "How RoadSOS satisfies it"
    for c, r in criteria:
        row = table.add_row().cells
        row[0].text = c
        row[1].text = r


def _architecture_section(doc: Document) -> None:
    add_heading(doc, "Architecture summary", level=1)
    add_para(doc,
             "The detailed 35-page technical reference lives at docs/TECHNICAL.pdf in the "
             "repository. The summary below captures the request flow and the four-tier "
             "offline strategy at a glance.",
             italic=True)

    add_heading(doc, "Request flow — GET /search", level=2)
    for line in [
        "1. Frontend issues GET /search?lat=…&lon=…&radius=5000.",
        "2. Backend phase 1: parallel fan-out — Nominatim reverse geocode + Overpass query (9 OSM categories, 3 mirrors, 12 s timeout each).",
        "3. Backend phase 2: parallel Google Places Nearby Search if a key is configured (does NOT wait for OSM; rankby=distance for nearest-first).",
        "4. Backend phase 3: merge + proximity dedupe (50 m radius), then phone enrichment via up to 3 Place Details calls within a 5 s budget.",
        "5. Backend response: JSON envelope with categorised contacts, country emergency numbers, landmark, and ISO country code.",
        "6. Frontend: render on Leaflet map (CartoDB Dark Matter tiles), populate Quick-Contacts dock, persist to 24 h localStorage cache.",
    ]:
        doc.add_paragraph(line, style="List Number")

    add_heading(doc, "Four-tier offline strategy", level=2)
    for line in [
        "Tier 1 — Live backend /search call (online path).",
        "Tier 2 — Service Worker + localStorage cache with 24 h TTL, keyed at ~1.1 km grid for high hit-rate.",
        "Tier 3 — Bundled JSON shipped with the PWA: 249 facilities across 196 countries.",
        "Tier 4 — Hardcoded mock as the final placeholder so the UI never shows an empty state.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "Crash detection signals", level=2)
    for line in [
        "Signal 1 — GPS velocity collapse: ≥ 25 km/h decaying to ≤ 5 km/h within 2 s.",
        "Signal 2 — Accelerometer spike: peak magnitude ≥ 3.5 G.",
        "Confirmation: both signals must occur within a 4 s alignment window.",
        "Cooldown: 12 s after a confirmed event to prevent retriggering.",
        "User override: a 10 s countdown lets the user PIN-cancel a false positive before auto-dispatch.",
    ]:
        doc.add_paragraph(line, style="List Bullet")


def _code_section(doc: Document) -> None:
    add_heading(doc, "Entire source code", level=1)
    add_para(doc,
             "Every file is rendered verbatim, grouped by directory. File paths are relative "
             "to the repository root. Binary assets (images, fonts, PDFs, lockfiles) are "
             "excluded. The full Git history is at https://github.com/Arthrevs/Roadproj.",
             italic=True)

    files_listed = 0
    files_missing: list[str] = []
    files_skipped_size: list[str] = []

    for group_title, files in CODE_GROUPS:
        add_heading(doc, group_title, level=2)
        for rel in files:
            path = ROOT / rel
            if not path.exists():
                files_missing.append(rel)
                add_para(doc, f"⚠ {rel} — not present in this revision", italic=True)
                continue
            try:
                content = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                content = path.read_text(encoding="utf-8", errors="replace")

            lang = LANG_BY_SUFFIX.get(path.suffix.lower(), "Text")
            add_heading(doc, f"{rel}  [{lang}, {len(content):,} chars]", level=3)
            add_code_block(doc, content)
            files_listed += 1

    add_page_break(doc)
    add_heading(doc, "Source index", level=2)
    add_para(doc, f"Files included verbatim: {files_listed}")
    if files_missing:
        add_para(doc, f"Files declared but missing on disk: {len(files_missing)}")
        for f in files_missing:
            doc.add_paragraph(f, style="List Bullet")


# ──────────────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────────────

def main() -> int:
    doc = Document()

    # Base style: 11pt Calibri.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RoadSOS")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x0B, 0x14, 0x24)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Hackathon submission package")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "\nNational Road Safety Hackathon 2026\n"
        "CoERS · Centre of Excellence in Road Safety · IIT Madras\n"
        "https://coers.iitm.ac.in/events/Hackathon/2026/rule_book/\n"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    contents = doc.add_paragraph()
    contents.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contents.add_run(
        "\nContents:\n"
        "1. Project overview\n"
        "2. Assumptions\n"
        "3. Software packages\n"
        "4. Architecture summary\n"
        "5. Entire source code\n"
    )
    run.font.size = Pt(12)

    add_page_break(doc)

    # ── Sections ───────────────────────────────────────────────────────────
    _overview_section(doc)
    add_page_break(doc)
    _assumptions_section(doc)
    add_page_break(doc)
    _packages_section(doc)
    add_page_break(doc)
    _architecture_section(doc)
    add_page_break(doc)
    _code_section(doc)

    out = ROOT / "docs" / "RoadSOS_Submission.docx"
    out.parent.mkdir(exist_ok=True)
    doc.save(out)
    size_kb = out.stat().st_size / 1024
    print(f"OK  wrote {out}  ({size_kb:,.1f} KB)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
